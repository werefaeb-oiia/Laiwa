import copy
from docx.table import Table
from docx.oxml import OxmlElement
from docx.shared import Inches, Pt  
from docx.oxml.ns import qn         

def set_thai_font(cell, text, font_name='TH SarabunPSK', font_size=14, style_name=None):
    cell.text = "" # เคลียร์ข้อความเก่าทิ้งก่อน
    if not str(text).strip():
        return # ถ้าไม่มีข้อความ ให้ปล่อยว่าง
        
    p = cell.paragraphs[0] if len(cell.paragraphs) > 0 else cell.add_paragraph()
    
    # --- [เพิ่มใหม่] ถ้ามีการส่งชื่อ Style มา ให้จับใส่ Paragraph เลย ---
    if style_name:
        try:
            p.style = style_name
        except Exception:
            pass # กันเหนียวไว้ เผื่อ Template Word ไม่มี Style นี้
            
    run = p.add_run(str(text))
    
    # กำหนดชื่อและขนาดฟอนต์ (ขนาดจะทับกับของ Heading เพื่อให้ได้ TH Sarabun สวยๆ)
    run.font.name = font_name
    run.font.size = Pt(font_size)
    
    # โค้ดบังคับให้ภาษาไทยเป็นฟอนต์ที่กำหนด
    r = run._element
    rPr = r.get_or_add_rPr()
    rFonts = rPr.get_or_add_rFonts()
    rFonts.set(qn('w:ascii'), font_name) 
    rFonts.set(qn('w:hAnsi'), font_name) 
    rFonts.set(qn('w:cs'), font_name)

def append_new_table(doc, previous_table, template_tbl_xml):
    p = OxmlElement('w:p')
    previous_table._element.addnext(p)
    new_tbl_elem = copy.deepcopy(template_tbl_xml)
    p.addnext(new_tbl_elem)
    new_table = Table(new_tbl_elem, doc)
    return new_table

def fill_table_data(table, row_data, images, log_callback):
    """
    รับ row_data (จาก Excel) มาเช็คและหยอดลงตารางทีละบรรทัด
    """
    is_success = False
    ts_no = str(row_data.get('*Test Script No.', '')).strip()
    
    for row in table.rows:
        if len(row.cells) >= 2:
            # คลีนข้อความคอลัมน์ซ้าย (ลบเว้นวรรค ลบ : และทำเป็นตัวเล็ก) เพื่อให้เปรียบเทียบง่ายขึ้น
            left_text = row.cells[0].text.lower().strip().replace(":", "")
            
            # --- เริ่มจับคู่ดึงข้อมูลจาก Excel ---
            if 'test script' in left_text:
                # สั่งให้ช่องนี้กลายเป็น Heading 3
                set_thai_font(row.cells[1], ts_no, style_name='Heading 3')
                
            elif 'test case name' in left_text:
                set_thai_font(row.cells[1], row_data.get('Test Case Name (Description)', ''))
                
            elif 'status' in left_text:
                pass # ข้ามไปก่อน
                
            elif 'pre-condition' in left_text or 'pre condition' in left_text:
                set_thai_font(row.cells[1], row_data.get('Pre-Condition', ''))
                
            elif 'step test' in left_text or 'test step' in left_text:
                set_thai_font(row.cells[1], row_data.get('Test Step Script', ''))
                
            elif 'expected result' in left_text:
                set_thai_font(row.cells[1], row_data.get('Expected Results', ''))
                
            # --- จัดการเรื่องรูปภาพ ---
            elif 'screenshot' in left_text:
                cell = row.cells[0]
                cell.text = "Screenshot:" # เคลียร์ข้อความเก่า
                
                if len(images) > 0:
                    for img in images:
                        p_img = cell.add_paragraph()
                        run = p_img.add_run()
                        run.add_picture(img, width=Inches(5.5))
                    is_success = True
                    log_callback(f"  -> Inserted {len(images)} images for {ts_no}")
                else:
                    log_callback(f"  -> Warning: No images found in folder: {ts_no}")
                    
    return is_success