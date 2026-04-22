import os
import glob
import copy
from datetime import datetime
from docx import Document
import pandas as pd # นำเข้าไลบรารีอ่าน Excel

from table_builder import append_new_table, fill_table_data

def generate_report(word_path, excel_path, base_folder, progress_callback, status_callback, log_callback, finish_callback):
    
    def log(msg):
        t = datetime.now().strftime("%H:%M:%S")
        log_callback(f"[{t}] {msg}")

    try:
        log("Starting report generation...")
        doc = Document(word_path)

        # -----------------------------
        # 1. อ่านข้อมูลจาก Excel
        # -----------------------------
        log("Reading Excel file...")
        try:
            # เพิ่ม header=1 เพื่อบอกให้ข้ามบรรทัดแรกไปอ่านแถวที่ 2 เป็นหัวตารางแทน
            try:
                df = pd.read_excel(excel_path, sheet_name='Setting', header=1)
            except ValueError:
                df = pd.read_excel(excel_path, header=1)
            
            # เคลียร์ช่องว่างที่แอบซ่อนอยู่ในชื่อคอลัมน์ (กันเหนียวไว้ก่อน)
            df.columns = df.columns.astype(str).str.strip()
            
            # เช็คความชัวร์ ถ้ายังหาไม่เจออีกให้ปริ้นบอกว่าเจอคอลัมน์อะไรบ้าง
            if '*Test Script No:' not in df.columns:
                log(f"Error: ไม่พบคอลัมน์ '*Test Script No:'")
                log(f"คอลัมน์ที่มองเห็นคือ: {list(df.columns)}")
                # ถ้าหัวตารางไปอยู่บรรทัดที่ 3 อาจจะต้องเปลี่ยนข้างบนเป็น header=2 ครับ
                raise ValueError("ไม่เจอชื่อคอลัมน์ที่ต้องการ เช็คใน Log ดูครับ")

            # เติมค่าว่าง (NaN) ให้เป็น String ว่างๆ
            df = df.fillna('')
            
            # ตัดแถวที่ช่อง *Test Script No: เป็นค่าว่างทิ้งไป
            df = df[df['*Test Script No:'] != '']
            
            # แปลงเป็น List ของ Dictionary
            excel_data = df.to_dict('records')
            total = len(excel_data)
            log(f"Found {total} test scripts in Excel.")
            
        except Exception as e:
            status_callback("Error reading Excel.", "error")
            log(f"Excel Error: {str(e)}")
            finish_callback()
            return

        if total == 0:
            status_callback("Error: No data in Excel.", "error")
            finish_callback()
            return

        if len(doc.tables) == 0:
            status_callback("Error: No table found in Word template.", "error")
            finish_callback()
            return

        template_table = doc.tables[0]
        template_tbl_xml = copy.deepcopy(template_table._tbl)
        previous_table = template_table
        
        success = 0

        # -----------------------------
        # 2. ลูปสร้างตารางตาม Excel
        # -----------------------------
        for i, row_data in enumerate(excel_data, start=1):
            ts_no = str(row_data.get('*Test Script No:', '')).strip()
            
            # ค้นหาโฟลเดอร์รูปภาพที่ชื่อตรงกับ Test Script No
            folder_full = os.path.join(base_folder, ts_no)
            images = []
            if os.path.isdir(folder_full):
                for ext in ["*.png", "*.jpg", "*.jpeg", "*.gif"]:
                    images.extend(glob.glob(os.path.join(folder_full, ext)))
                images.sort()

            # สร้างตาราง
            if i == 1:
                current_table = template_table
                log(f"Using existing template table for {ts_no}")
            else:
                current_table = append_new_table(doc, previous_table, template_tbl_xml)
                previous_table = current_table
                log(f"Generated new table for {ts_no}")

            # ส่งข้อมูลจาก Excel และรูปภาพไปยัดลงตาราง
            if fill_table_data(current_table, row_data, images, log):
                success += 1

            progress_callback(i / total if total > 0 else 1)
            status_callback(f"Processing: {ts_no} ({i}/{total})", "normal")

        # Save ไฟล์
        now = datetime.now().strftime("%Y%m%d_%H%M%S")
        new_path = word_path.replace(".docx", f"_{now}.docx")
        log(f"Saving output file: {os.path.basename(new_path)}...")
        doc.save(new_path)

        status_callback(f"Done! Created {total} tables, {success} have images.", "success")
        log(f"Process completed successfully!")
        
    except Exception as e:
        status_callback("Error occurred (Check log)", "error")
        log(f"CRITICAL ERROR: {str(e)}")
        
    finally:
        finish_callback()